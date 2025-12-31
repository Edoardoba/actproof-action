"""
Generatore Report Legali
Genera PDF/Docx ufficiali pronti per ispezione autorità (AgID/ACN)
"""

from pathlib import Path
from typing import Optional, Dict, Any
from datetime import datetime
from actproof.fairness.auditor import BiasReport
from actproof.compliance.requirements import ComplianceResult, TechnicalDocumentation


class LegalReportGenerator:
    """
    Genera report legali in formato PDF/Docx
    Pronti per ispezione autorità (AgID/ACN)
    """
    
    def __init__(self):
        """Inizializza generatore report"""
        self._check_dependencies()
    
    def _check_dependencies(self):
        """Verifica dipendenze per generazione report"""
        self.reportlab_available = False
        self.docx_available = False
        
        try:
            import reportlab
            from reportlab.lib.pagesizes import A4
            from reportlab.lib import colors
            self.reportlab_available = True
        except ImportError:
            print("⚠️  reportlab non disponibile. Installa con: pip install reportlab")
        
        try:
            from docx import Document
            self.docx_available = True
        except ImportError:
            print("⚠️  python-docx non disponibile. Installa con: pip install python-docx")
    
    def generate_pdf_report(
        self,
        bias_report: BiasReport,
        compliance_result: Optional[ComplianceResult] = None,
        technical_doc: Optional[TechnicalDocumentation] = None,
        output_path: Optional[Path] = None,
    ) -> Path:
        """
        Genera report PDF ufficiale
        
        Args:
            bias_report: Report bias e fairness
            compliance_result: Risultato conformità (opzionale)
            technical_doc: Documentazione tecnica (opzionale)
            output_path: Percorso output (default: reports/bias_report_{timestamp}.pdf)
        
        Returns:
            Percorso del file generato
        """
        if not self.reportlab_available:
            raise ImportError("reportlab non disponibile. Installa con: pip install reportlab")
        
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
        from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
        
        if output_path is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = Path(f"reports/bias_report_{timestamp}.pdf")
        
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Crea documento PDF
        doc = SimpleDocTemplate(str(output_path), pagesize=A4)
        story = []
        styles = getSampleStyleSheet()
        
        # Titolo
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=colors.HexColor('#1a237e'),
            spaceAfter=30,
            alignment=TA_CENTER,
        )
        story.append(Paragraph("Report di Fairness e Bias", title_style))
        story.append(Paragraph(f"Sistema: {bias_report.model_name}", styles['Normal']))
        story.append(Paragraph(f"Data: {bias_report.evaluated_at.strftime('%d/%m/%Y %H:%M')}", styles['Normal']))
        story.append(Spacer(1, 0.5*cm))
        
        # Sezione: Riepilogo Conformità
        story.append(Paragraph("Riepilogo Conformità", styles['Heading2']))
        compliance_status = "✅ CONFORME" if bias_report.overall_compliant else "❌ NON CONFORME"
        story.append(Paragraph(f"<b>Stato Complessivo:</b> {compliance_status}", styles['Normal']))
        story.append(Spacer(1, 0.3*cm))
        
        # Sezione: Metriche Fairness
        story.append(Paragraph("Metriche di Fairness per Attributo Protetto", styles['Heading2']))
        
        for attr_name, metrics in bias_report.fairness_metrics.items():
            story.append(Paragraph(f"<b>Attributo Protetto: {attr_name}</b>", styles['Heading3']))
            
            # Tabella metriche
            data = [
                ['Metrica', 'Valore', 'Conforme'],
                ['Demographic Parity Difference (DPD)', f"{metrics.demographic_parity_difference:.4f}", 
                 "✅" if metrics.compliant_dpd else "❌"],
                ['Equalized Odds Difference (EOD)', f"{metrics.equalized_odds_difference:.4f}",
                 "✅" if metrics.compliant_eod else "❌"],
                ['False Positive Rate Difference', f"{metrics.false_positive_rate_difference:.4f}", "-"],
                ['False Negative Rate Difference', f"{metrics.false_negative_rate_difference:.4f}", "-"],
            ]
            
            table = Table(data, colWidths=[8*cm, 4*cm, 2*cm])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ]))
            story.append(table)
            story.append(Spacer(1, 0.3*cm))
            
            # Metriche per gruppo
            if metrics.group_metrics:
                story.append(Paragraph("<b>Metriche per Gruppo:</b>", styles['Normal']))
                group_data = [['Gruppo', 'FPR', 'FNR', 'Selection Rate']]
                for group, group_metrics in metrics.group_metrics.items():
                    group_data.append([
                        group,
                        f"{group_metrics['fpr']:.4f}",
                        f"{group_metrics['fnr']:.4f}",
                        f"{group_metrics['selection_rate']:.4f}",
                    ])
                
                group_table = Table(group_data, colWidths=[3*cm, 3*cm, 3*cm, 3*cm])
                group_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ]))
                story.append(group_table)
                story.append(Spacer(1, 0.3*cm))
        
        # Sezione: Bias Critici
        if bias_report.critical_biases:
            story.append(Paragraph("Bias Critici Identificati", styles['Heading2']))
            for bias in bias_report.critical_biases:
                story.append(Paragraph(f"• {bias}", styles['Normal']))
            story.append(Spacer(1, 0.3*cm))
        
        # Sezione: Raccomandazioni
        if bias_report.recommendations:
            story.append(Paragraph("Raccomandazioni", styles['Heading2']))
            for rec in bias_report.recommendations:
                story.append(Paragraph(rec, styles['Normal']))
            story.append(Spacer(1, 0.3*cm))
        
        # Sezione: Informazioni Conformità (se disponibile)
        if compliance_result:
            story.append(PageBreak())
            story.append(Paragraph("Informazioni Conformità EU AI Act", styles['Heading2']))
            story.append(Paragraph(f"<b>Sistema ID:</b> {compliance_result.system_id}", styles['Normal']))
            story.append(Paragraph(f"<b>Livello di Rischio:</b> {compliance_result.risk_level.value}", styles['Normal']))
            story.append(Paragraph(f"<b>Score Conformità:</b> {compliance_result.requirements_check.compliance_score:.2%}", styles['Normal']))
            story.append(Spacer(1, 0.3*cm))
        
        # Footer con informazioni legali
        story.append(Spacer(1, 1*cm))
        footer_style = ParagraphStyle(
            'Footer',
            parent=styles['Normal'],
            fontSize=8,
            textColor=colors.grey,
            alignment=TA_CENTER,
        )
        story.append(Paragraph(
            "Questo report è stato generato automaticamente da ActProof.ai per conformità EU AI Act (Regolamento UE 2024/1689). "
            "Per domande o chiarimenti, contattare le autorità competenti (AgID/ACN).",
            footer_style
        ))
        
        # Genera PDF
        doc.build(story)
        return output_path
    
    def generate_docx_report(
        self,
        bias_report: BiasReport,
        compliance_result: Optional[ComplianceResult] = None,
        technical_doc: Optional[TechnicalDocumentation] = None,
        output_path: Optional[Path] = None,
    ) -> Path:
        """
        Genera report Docx ufficiale
        
        Args:
            bias_report: Report bias e fairness
            compliance_result: Risultato conformità (opzionale)
            technical_doc: Documentazione tecnica (opzionale)
            output_path: Percorso output
        
        Returns:
            Percorso del file generato
        """
        if not self.docx_available:
            raise ImportError("python-docx non disponibile. Installa con: pip install python-docx")
        
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        if output_path is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = Path(f"reports/bias_report_{timestamp}.docx")
        
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Crea documento Word
        doc = Document()
        
        # Titolo
        title = doc.add_heading('Report di Fairness e Bias', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Informazioni sistema
        doc.add_paragraph(f'Sistema: {bias_report.model_name}')
        doc.add_paragraph(f'Data: {bias_report.evaluated_at.strftime("%d/%m/%Y %H:%M")}')
        doc.add_paragraph()
        
        # Riepilogo Conformità
        doc.add_heading('Riepilogo Conformità', level=1)
        compliance_status = "✅ CONFORME" if bias_report.overall_compliant else "❌ NON CONFORME"
        doc.add_paragraph(f'Stato Complessivo: {compliance_status}')
        doc.add_paragraph()
        
        # Metriche Fairness
        doc.add_heading('Metriche di Fairness per Attributo Protetto', level=1)
        
        for attr_name, metrics in bias_report.fairness_metrics.items():
            doc.add_heading(f'Attributo Protetto: {attr_name}', level=2)
            
            # Tabella metriche
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Light Grid Accent 1'
            
            # Header
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Metrica'
            header_cells[1].text = 'Valore'
            header_cells[2].text = 'Conforme'
            
            # Righe dati
            rows_data = [
                ('Demographic Parity Difference (DPD)', f"{metrics.demographic_parity_difference:.4f}",
                 "✅" if metrics.compliant_dpd else "❌"),
                ('Equalized Odds Difference (EOD)', f"{metrics.equalized_odds_difference:.4f}",
                 "✅" if metrics.compliant_eod else "❌"),
                ('False Positive Rate Difference', f"{metrics.false_positive_rate_difference:.4f}", "-"),
                ('False Negative Rate Difference', f"{metrics.false_negative_rate_difference:.4f}", "-"),
            ]
            
            for row_data in rows_data:
                row_cells = table.add_row().cells
                row_cells[0].text = row_data[0]
                row_cells[1].text = row_data[1]
                row_cells[2].text = row_data[2]
            
            doc.add_paragraph()
            
            # Metriche per gruppo
            if metrics.group_metrics:
                doc.add_paragraph('Metriche per Gruppo:', style='Heading 3')
                group_table = doc.add_table(rows=1, cols=4)
                group_table.style = 'Light Grid Accent 1'
                
                group_header = group_table.rows[0].cells
                group_header[0].text = 'Gruppo'
                group_header[1].text = 'FPR'
                group_header[2].text = 'FNR'
                group_header[3].text = 'Selection Rate'
                
                for group, group_metrics in metrics.group_metrics.items():
                    group_row = group_table.add_row().cells
                    group_row[0].text = str(group)
                    group_row[1].text = f"{group_metrics['fpr']:.4f}"
                    group_row[2].text = f"{group_metrics['fnr']:.4f}"
                    group_row[3].text = f"{group_metrics['selection_rate']:.4f}"
                
                doc.add_paragraph()
        
        # Bias Critici
        if bias_report.critical_biases:
            doc.add_heading('Bias Critici Identificati', level=1)
            for bias in bias_report.critical_biases:
                doc.add_paragraph(bias, style='List Bullet')
            doc.add_paragraph()
        
        # Raccomandazioni
        if bias_report.recommendations:
            doc.add_heading('Raccomandazioni', level=1)
            for rec in bias_report.recommendations:
                doc.add_paragraph(rec)
            doc.add_paragraph()
        
        # Informazioni Conformità
        if compliance_result:
            doc.add_page_break()
            doc.add_heading('Informazioni Conformità EU AI Act', level=1)
            doc.add_paragraph(f'Sistema ID: {compliance_result.system_id}')
            doc.add_paragraph(f'Livello di Rischio: {compliance_result.risk_level.value}')
            doc.add_paragraph(f'Score Conformità: {compliance_result.requirements_check.compliance_score:.2%}')
            doc.add_paragraph()
        
        # Footer
        footer_para = doc.add_paragraph()
        footer_para.add_run(
            'Questo report è stato generato automaticamente da ActProof.ai per conformità EU AI Act '
            '(Regolamento UE 2024/1689). Per domande o chiarimenti, contattare le autorità competenti (AgID/ACN).'
        ).font.size = Pt(8)
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Salva documento
        doc.save(str(output_path))
        return output_path
